"""Raw text extraction for the docx-driven page types (category/blog) - see
app.planner.image_planner.DOCX_DRIVEN_PAGE_TYPES. Deliberately just plain text, not structured
JSON: these page types have no facts schema behind them, so the Prompt Generator grounds itself
in the whole document rather than a subset of named fields."""

import io

from docx import Document


class DocxExtractionError(ValueError):
    pass


def extract_text_from_docx(file_bytes: bytes) -> str:
    try:
        document = Document(io.BytesIO(file_bytes))
    except Exception as exc:  # noqa: BLE001 - python-docx raises assorted errors for non-docx input
        raise DocxExtractionError(f"Could not read the uploaded file as a .docx document: {exc}") from exc

    paragraphs = [p.text for p in document.paragraphs if p.text.strip()]
    # Tables often carry facts that never appear in plain paragraphs (fee tables, comparisons) -
    # include their cell text too so the Prompt Generator isn't blind to them.
    for table in document.tables:
        for row in table.rows:
            cells_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells_text:
                paragraphs.append(" | ".join(cells_text))

    text = "\n".join(paragraphs)
    if not text.strip():
        raise DocxExtractionError("No extractable text found in the uploaded document")
    return text
